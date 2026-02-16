"""Word操作工具类"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK_TYPE
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from typing import Any, List, Dict, Optional, Tuple
import os


class WordUtils:
    """
    Word操作工具类，提供Word文档的读写功能
    """

    @staticmethod
    def create_document() -> Document:
        """
        创建新的Word文档
        
        Returns:
            Document对象
        """
        return Document()

    @staticmethod
    def open_document(file_path: str) -> Document:
        """
        打开现有的Word文档
        
        Args:
            file_path: Word文档路径
            
        Returns:
            Document对象
            
        Raises:
            FileNotFoundError: 文件不存在
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        return Document(file_path)

    @staticmethod
    def save_document(document: Document, file_path: str) -> None:
        """
        保存Word文档
        
        Args:
            document: Document对象
            file_path: 保存路径
        """
        # 确保目录存在
        os.makedirs(os.path.dirname(file_path) if os.path.dirname(file_path) else '.', exist_ok=True)
        document.save(file_path)

    @staticmethod
    def add_heading(document: Document, text: str, level: int = 1) -> Any:
        """
        添加标题
        
        Args:
            document: Document对象
            text: 标题文本
            level: 标题级别，1-9
            
        Returns:
            标题对象
        """
        return document.add_heading(text, level=level)

    @staticmethod
    def add_paragraph(document: Document, text: str, style: Optional[str] = None) -> Any:
        """
        添加段落
        
        Args:
            document: Document对象
            text: 段落文本
            style: 段落样式
            
        Returns:
            段落对象
        """
        return document.add_paragraph(text, style=style)

    @staticmethod
    def add_table(document: Document, rows: int, cols: int, style: Optional[str] = None) -> Any:
        """
        添加表格
        
        Args:
            document: Document对象
            rows: 行数
            cols: 列数
            style: 表格样式
            
        Returns:
            表格对象
        """
        table = document.add_table(rows=rows, cols=cols)
        if style:
            table.style = style
        return table

    @staticmethod
    def add_picture(document: Document, image_path: str, width: Optional[float] = None) -> Any:
        """
        添加图片
        
        Args:
            document: Document对象
            image_path: 图片路径
            width: 图片宽度（英寸）
            
        Returns:
            图片对象
        """
        if width:
            return document.add_picture(image_path, width=Inches(width))
        else:
            return document.add_picture(image_path)

    @staticmethod
    def add_page_break(document: Document) -> None:
        """
        添加分页符
        
        Args:
            document: Document对象
        """
        paragraph = document.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK_TYPE.PAGE)

    @staticmethod
    def set_paragraph_alignment(paragraph: Any, alignment: str) -> None:
        """
        设置段落对齐方式
        
        Args:
            paragraph: 段落对象
            alignment: 对齐方式，可选值：left, center, right, justify
        """
        alignment_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        }
        if alignment in alignment_map:
            paragraph.alignment = alignment_map[alignment]

    @staticmethod
    def set_font_style(run: Any, font_name: str = 'Arial', font_size: int = 12, 
                     bold: bool = False, italic: bool = False, 
                     color: Optional[Tuple[int, int, int]] = None) -> None:
        """
        设置字体样式
        
        Args:
            run: 文本运行对象
            font_name: 字体名称
            font_size: 字体大小
            bold: 是否加粗
            italic: 是否斜体
            color: 字体颜色，RGB值 (r, g, b)
        """
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)

    @staticmethod
    def create_table_from_data(document: Document, data: List[List[Any]], 
                             headers: Optional[List[str]] = None, 
                             style: Optional[str] = None) -> Any:
        """
        从数据创建表格
        
        Args:
            document: Document对象
            data: 表格数据，二维列表
            headers: 表头列表
            style: 表格样式
            
        Returns:
            表格对象
        """
        # 计算行数和列数
        rows = len(data) + (1 if headers else 0)
        cols = max(len(row) for row in data) if data else 0
        if headers:
            cols = max(cols, len(headers))
        
        # 创建表格
        table = WordUtils.add_table(document, rows, cols, style)
        
        # 添加表头
        if headers:
            for i, header in enumerate(headers):
                if i < cols:
                    table.cell(0, i).text = str(header)
        
        # 添加数据
        start_row = 1 if headers else 0
        for i, row_data in enumerate(data):
            for j, cell_data in enumerate(row_data):
                if j < cols:
                    table.cell(start_row + i, j).text = str(cell_data)
        
        return table

    @staticmethod
    def read_document(file_path: str) -> Dict[str, Any]:
        """
        读取Word文档内容
        
        Args:
            file_path: Word文档路径
            
        Returns:
            文档内容字典，包含标题、段落和表格
        """
        document = WordUtils.open_document(file_path)
        content = {
            'headings': [],
            'paragraphs': [],
            'tables': []
        }
        
        # 读取标题和段落
        for para in document.paragraphs:
            if para.style.name.startswith('Heading'):
                content['headings'].append({
                    'text': para.text,
                    'level': int(para.style.name.split(' ')[-1])
                })
            else:
                content['paragraphs'].append(para.text)
        
        # 读取表格
        for table in document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                table_data.append(row_data)
            content['tables'].append(table_data)
        
        return content

    @staticmethod
    def set_page_orientation(document: Document, orientation: str = 'portrait') -> None:
        """
        设置页面方向
        
        Args:
            document: Document对象
            orientation: 页面方向，可选值：portrait（纵向）, landscape（横向）
        """
        section = document.sections[0]
        if orientation == 'landscape':
            section.orientation = WD_ORIENTATION.LANDSCAPE
            # 交换页面宽度和高度
            width, height = section.page_width, section.page_height
            section.page_width = height
            section.page_height = width
        else:  # portrait
            section.orientation = WD_ORIENTATION.PORTRAIT

    @staticmethod
    def set_page_margins(document: Document, top: float = 1.0, bottom: float = 1.0, 
                        left: float = 1.0, right: float = 1.0) -> None:
        """
        设置页面边距
        
        Args:
            document: Document对象
            top: 上边距（英寸）
            bottom: 下边距（英寸）
            left: 左边距（英寸）
            right: 右边距（英寸）
        """
        section = document.sections[0]
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

    @staticmethod
    def add_hyperlink(document: Document, text: str, url: str) -> Any:
        """
        添加超链接
        
        Args:
            document: Document对象
            text: 链接文本
            url: 链接地址
            
        Returns:
            段落对象
        """
        paragraph = document.add_paragraph()
        run = paragraph.add_run(text)
        
        # 添加超链接
        # 注意：python-docx不直接支持添加超链接，需要使用XML
        part = document.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        
        # 获取运行的XML
        r = run._r
        rPr = r.get_or_add_rPr()
        
        # 创建超链接元素
        hyperlink = r.add_hyperlink()
        hyperlink.set('r:id', r_id)
        
        # 移动文本到超链接中
        for child in list(r):
            if child.tag != rPr.tag:
                hyperlink.append(child)
        
        return paragraph

    @staticmethod
    def create_simple_report(file_path: str, title: str, content: Dict[str, Any]) -> None:
        """
        创建简单的Word报告
        
        Args:
            file_path: 保存路径
            title: 报告标题
            content: 报告内容，包含sections（章节）和tables（表格）
        """
        # 创建文档
        doc = WordUtils.create_document()
        
        # 添加标题
        WordUtils.add_heading(doc, title, level=0)
        WordUtils.add_page_break(doc)
        
        # 添加章节
        if 'sections' in content:
            for section in content['sections']:
                if 'title' in section:
                    WordUtils.add_heading(doc, section['title'], level=1)
                if 'content' in section:
                    WordUtils.add_paragraph(doc, section['content'])
                WordUtils.add_paragraph(doc, '')  # 空行
        
        # 添加表格
        if 'tables' in content:
            for table_data in content['tables']:
                if 'title' in table_data:
                    WordUtils.add_heading(doc, table_data['title'], level=2)
                if 'headers' in table_data and 'data' in table_data:
                    WordUtils.create_table_from_data(
                        doc, table_data['data'], table_data['headers']
                    )
                WordUtils.add_paragraph(doc, '')  # 空行
        
        # 保存文档
        WordUtils.save_document(doc, file_path)